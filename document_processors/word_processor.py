"""
Word文档处理器
处理DOCX和DOC格式文档的文本提取和元数据提取
"""

import logging
from typing import List, Dict, Any
import os

from .base_processor import BaseDocumentProcessor

# 日志配置
logger = logging.getLogger(__name__)

class WordProcessor(BaseDocumentProcessor):
    """Word文档处理器"""
    
    def __init__(self):
        """初始化Word处理器"""
        super().__init__()
    
    def can_process(self, file_path: str) -> bool:
        """判断是否可以处理Word文件"""
        return file_path.lower().endswith(('.docx', '.doc'))
    
    def get_supported_extensions(self) -> List[str]:
        """获取支持的文件扩展名"""
        return ['.docx', '.doc']
    
    def extract_text(self, file_path: str) -> str:
        """
        从Word文件中提取文本
        
        Args:
            file_path: Word文件路径
        
        Returns:
            提取的文本内容
        """
        try:
            # 优先使用python-docx处理.docx文件
            if file_path.lower().endswith('.docx'):
                return self._extract_from_docx(file_path)
            else:
                # 对于.doc文件，使用docx2txt
                return self._extract_from_doc(file_path)
                
        except Exception as e:
            logger.error(f"从Word文档提取文本失败: {file_path} - {e}")
            return ""
    
    def _extract_from_docx(self, file_path: str) -> str:
        """从DOCX文件提取文本（使用python-docx）"""
        try:
            from docx import Document
            doc = Document(file_path)
            
            content = []
            
            # 提取标题和段落
            for para in doc.paragraphs:
                if para.text.strip():
                    # 检查是否是标题
                    if para.style.name.startswith('Heading'):
                        content.append(f"\n=== {para.text.strip()} ===\n")
                    else:
                        content.append(para.text.strip())
            
            # 提取表格内容
            if doc.tables:
                content.append("\n" + "="*50)
                content.append("文档表格内容:")
                content.append("="*50)
                
                for i, table in enumerate(doc.tables):
                    content.append(f"\n--- 表格 {i+1} ---")
                    for row in table.rows:
                        row_text = " | ".join(cell.text.strip() for cell in row.cells if cell.text.strip())
                        if row_text:
                            content.append(row_text)
            
            return "\n".join(content)
            
        except Exception as e:
            logger.error(f"使用python-docx提取文本失败: {file_path} - {e}")
            # 降级到docx2txt
            return self._extract_with_docx2txt(file_path)
    
    def _extract_from_doc(self, file_path: str) -> str:
        """从DOC文件提取文本（使用docx2txt）"""
        return self._extract_with_docx2txt(file_path)
    
    def _extract_with_docx2txt(self, file_path: str) -> str:
        """使用docx2txt提取文本"""
        try:
            import docx2txt
            text = docx2txt.process(file_path)
            return text.strip() if text else ""
        except Exception as e:
            logger.error(f"使用docx2txt提取文本失败: {file_path} - {e}")
            return ""
    
    def extract_metadata(self, file_path: str) -> Dict[str, Any]:
        """
        提取Word文档元数据
        
        Args:
            file_path: Word文件路径
            
        Returns:
            Word元数据字典
        """
        try:
            metadata_dict = {
                'document_type': 'Word',
                'file_extension': os.path.splitext(file_path)[1].lower()
            }
            
            # 尝试从DOCX文件提取详细元数据
            if file_path.lower().endswith('.docx'):
                try:
                    from docx import Document
                    doc = Document(file_path)
                    
                    # 获取文档核心属性
                    core_props = doc.core_properties
                    metadata_dict.update({
                        'title': core_props.title or '',
                        'author': core_props.author or '',
                        'subject': core_props.subject or '',
                        'keywords': core_props.keywords or '',
                        'comments': core_props.comments or '',
                        'last_modified_by': core_props.last_modified_by or '',
                        'created': str(core_props.created) if core_props.created else '',
                        'modified': str(core_props.modified) if core_props.modified else '',
                        'category': core_props.category or '',
                        'language': core_props.language or ''
                    })
                    
                    # 统计文档结构
                    paragraph_count = len([p for p in doc.paragraphs if p.text.strip()])
                    table_count = len(doc.tables)
                    
                    # 统计标题层级
                    heading_stats = {}
                    for para in doc.paragraphs:
                        if para.style.name.startswith('Heading'):
                            heading_level = para.style.name
                            heading_stats[heading_level] = heading_stats.get(heading_level, 0) + 1
                    
                    metadata_dict.update({
                        'paragraph_count': paragraph_count,
                        'table_count': table_count,
                        'heading_stats': heading_stats,
                        'has_tables': table_count > 0,
                        'has_headings': len(heading_stats) > 0
                    })
                    
                except Exception as e:
                    logger.warning(f"提取DOCX详细元数据失败: {file_path} - {e}")
            
            return metadata_dict
            
        except Exception as e:
            logger.error(f"提取Word元数据失败: {file_path} - {e}")
            return {
                'document_type': 'Word',
                'error': str(e)
            } 